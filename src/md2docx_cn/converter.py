"""Convert a focused Markdown subset into a polished Chinese DOCX file."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


@dataclass(frozen=True, slots=True)
class Block:
    """A parsed Markdown block used by the document writer."""

    kind: str
    text: str
    level: int = 0


HEADING_RE = re.compile(r"^(#{1,3})\s+(.+)$")
BULLET_RE = re.compile(r"^[-*+]\s+(.+)$")
NUMBERED_RE = re.compile(r"^\d+[.)]\s+(.+)$")
QUOTE_RE = re.compile(r"^>\s*(.+)$")


def parse_markdown(markdown: str) -> list[Block]:
    """Parse headings, bullets, and paragraphs from a Markdown article."""
    blocks: list[Block] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            blocks.append(Block("paragraph", " ".join(paragraph_lines)))
            paragraph_lines.clear()

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            flush_paragraph()
            blocks.append(
                Block("heading", heading_match.group(2), len(heading_match.group(1)))
            )
            continue

        bullet_match = BULLET_RE.match(line)
        if bullet_match:
            flush_paragraph()
            blocks.append(Block("bullet", bullet_match.group(1)))
            continue

        numbered_match = NUMBERED_RE.match(line)
        if numbered_match:
            flush_paragraph()
            blocks.append(Block("numbered", numbered_match.group(1)))
            continue

        quote_match = QUOTE_RE.match(line)
        if quote_match:
            flush_paragraph()
            blocks.append(Block("quote", quote_match.group(1)))
            continue

        paragraph_lines.append(line)

    flush_paragraph()
    return blocks


def _set_east_asia_font(style, font_name: str, size: float) -> None:
    style.font.name = font_name
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    style.font.size = Pt(size)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    _set_east_asia_font(normal, "SimSun", 12)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, 18, 10, "2E74B5"),
        "Heading 2": (13, 14, 7, "2E74B5"),
        "Heading 3": (12, 10, 5, "1F4D78"),
    }
    for style_name, (size, before, after, color) in heading_tokens.items():
        style = document.styles[style_name]
        _set_east_asia_font(style, "Microsoft YaHei", size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    bullet_style = document.styles["List Bullet"]
    _set_east_asia_font(bullet_style, "SimSun", 12)
    bullet_style.paragraph_format.left_indent = Inches(0.375)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.188)
    bullet_style.paragraph_format.space_after = Pt(4)
    bullet_style.paragraph_format.line_spacing = 1.25

    numbered_style = document.styles["List Number"]
    _set_east_asia_font(numbered_style, "SimSun", 12)
    numbered_style.paragraph_format.left_indent = Inches(0.375)
    numbered_style.paragraph_format.first_line_indent = Inches(-0.188)
    numbered_style.paragraph_format.space_after = Pt(4)
    numbered_style.paragraph_format.line_spacing = 1.25

    quote_style = document.styles["Quote"]
    _set_east_asia_font(quote_style, "KaiTi", 11)
    quote_style.font.color.rgb = RGBColor.from_string("555555")
    quote_style.paragraph_format.left_indent = Inches(0.375)
    quote_style.paragraph_format.right_indent = Inches(0.375)
    quote_style.paragraph_format.space_before = Pt(4)
    quote_style.paragraph_format.space_after = Pt(6)


def _add_blocks(document: Document, blocks: list[Block]) -> None:
    for block in blocks:
        if block.kind == "heading":
            paragraph = document.add_paragraph(
                block.text, style=f"Heading {min(block.level, 3)}"
            )
            if block.level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif block.kind == "bullet":
            document.add_paragraph(block.text, style="List Bullet")
        elif block.kind == "numbered":
            document.add_paragraph(block.text, style="List Number")
        elif block.kind == "quote":
            document.add_paragraph(block.text, style="Quote")
        else:
            paragraph = document.add_paragraph(block.text, style="Normal")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Pt(24)


def convert_markdown(
    source: str | Path,
    output: str | Path,
    *,
    author: str = "",
) -> Path:
    """Convert a UTF-8 Markdown file into a Chinese article DOCX."""
    source_path = Path(source)
    output_path = Path(output)
    if not source_path.is_file():
        raise FileNotFoundError(f"Markdown source not found: {source_path}")

    blocks = parse_markdown(source_path.read_text(encoding="utf-8"))
    document = Document()
    _configure_document(document)
    _add_blocks(document, blocks)

    if blocks and blocks[0].kind == "heading":
        document.core_properties.title = blocks[0].text
    document.core_properties.author = author
    document.core_properties.subject = "Generated by md2docx-cn"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path
